'''
Created on Feb 2, 2021
@author: Jiri Martinek at jimar@kiv.zcu.cz
'''

import os
import tempfile
from docx import Document
from docx.shared import Pt
import config
from docx2pdf import convert
from pdf2image import convert_from_path
from subprocess import  Popen
import sys
import shutil
import time

LIBRE_OFFICE = r"loffice"





class DocumentWriter:

    def __init__(self):
        self.document = Document()
        self.dataset_name = config.DATASET_NAME

    '''
    Linux based function for converting .docx to PDF 
    @:param input docx document
    @:param out folder   
    '''
    def convert_to_pdf(self, input_docx, out_folder):
        p = Popen([LIBRE_OFFICE, '--headless', '--convert-to', 'pdf', '--outdir',
                   out_folder, input_docx])
        print([LIBRE_OFFICE, '--convert-to', 'pdf', input_docx])
        p.communicate()

    '''
        Stores individual pages into JPEG format 
        @:param list of images imgs    
        '''
    def convert_to_images(self):
        print("Converting to pdf")
        if "win" in sys.platform:
            convert(self.dataset_name + ".docx")
        else:
            sample_doc = self.dataset_name + ".docx"
            out_folder = '.'
            self.convert_to_pdf(sample_doc, out_folder)
        time.sleep(2)
        with tempfile.TemporaryDirectory() as path:
            images = convert_from_path(self.dataset_name + ".pdf", output_folder=path)
            self.save_images(images)

    '''
    Stores individual pages into JPEG format 
    @:param list of images imgs    
    '''
    def save_images(self, imgs):
        dataset_folder = self.dataset_name
        if os.path.exists(dataset_folder):
            shutil.rmtree(dataset_folder)
            os.makedirs(dataset_folder)
        else:
            os.makedirs(dataset_folder)

        print("converting to JPG files")
        counter = 0
        for image in imgs:
            counter += + 1
            # print("image" + str(counter))
            path = os.path.join(self.dataset_name, "image" + str(counter) + ".jpg")
            image.save(path)
        print(str(counter)+ " images have been saved...")

    '''
    Creation of the dataset (from the input conll file)
    @:param input_connl_file    
    '''
    def create_dataset(self, input_conll_file):

        with open(input_conll_file, mode="r", encoding="utf8") as fr:
            conll_data = fr.readlines()

        style = self.document.styles['Normal']
        font = style.font
        font.name = config.FONT
        font.size = Pt(config.FONT_SIZE)
        text = ""

        counter = 0

        for line in conll_data:
            line_elements = line.split()
            if len(line_elements) == 0:
                paragraph = self.document.add_paragraph(text)
                paragraph.style = style
                text = ""
            elif line_elements[0] == ";":
                paragraph = self.document.add_paragraph(text)
                paragraph.style = style
                self.document.add_page_break()
                counter += 1
                if counter % 100 == 0:
                    print("Processed: ", counter + 1)
                text = ""
            else:
                text = text + " " + line_elements[0]

        self.document.save(self.dataset_name + ".docx")

